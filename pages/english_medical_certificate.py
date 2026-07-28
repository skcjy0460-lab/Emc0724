# -*- coding: utf-8 -*-
"""
해외 출국용 영문 진단서 자동 변환 도구
====================================

한글 진단서(스캔/사진)를 업로드하면 Gemini Vision으로 항목을 추출하고,
용도(Fit to Fly / 보험 청구용 / 일반 진단서)에 맞는 영문 서식에 데이터를
채워 Word(.docx) 초안을 생성합니다.

⚠️ 중요: 이 도구가 생성하는 문서는 "영문 번역 초안"입니다.
   반드시 담당 의사의 검토 및 서명, 병원 직인 날인 후 제출해야
   법적 효력이 있는 공식 서류가 됩니다. (이 안내 문구는 모든 출력물에
   자동으로 포함됩니다.)

배포 방법은 기존 프로젝트와 동일합니다:
  - requirements.txt 에 아래 패키지 추가
      streamlit
      google-genai
      python-docx
      pandas
  - GitHub 저장소에 이 파일 + diagnosis_mapping.csv 함께 업로드
  - Streamlit Cloud에서 멀티페이지 앱의 pages/ 폴더에 넣으면 됩니다.
  - Gemini API 키는 Streamlit Cloud의 Secrets에 GEMINI_API_KEY 로 등록하세요.
"""

import os
import io
import json
import base64
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------------------------
# 0. 기본 설정
# ---------------------------------------------------------------------------

st.set_page_config(page_title="영문 진단서 변환 도구", page_icon="🩺", layout="wide")

MAPPING_CSV_PATH = os.path.join(os.path.dirname(__file__), "diagnosis_mapping.csv")

PURPOSE_OPTIONS = {
    "fit_to_fly": "✈️ Fit to Fly (비행 적합성 확인서)",
    "insurance": "📄 해외 보험 청구용 진단서",
    "general": "🏥 일반 영문 진단서 (비자/제출용)",
}

DISCLAIMER_TEXT = (
    "This document is an English-language draft translated with the assistance "
    "of an automated tool from the original Korean medical certificate. "
    "It is NOT valid until reviewed, corrected if necessary, and signed/stamped "
    "by the issuing physician and the hospital. Please have the attending "
    "physician verify all contents before this document is submitted to any "
    "airline, embassy, insurer, or other institution."
)


# ---------------------------------------------------------------------------
# 1. 병명 매핑 테이블 로드
# ---------------------------------------------------------------------------

@st.cache_data
def load_mapping_table() -> pd.DataFrame:
    empty_df = pd.DataFrame(columns=["kcd_code", "icd10_code", "korean_name", "english_name"])
    if not os.path.exists(MAPPING_CSV_PATH):
        return empty_df
    try:
        return pd.read_csv(MAPPING_CSV_PATH, dtype=str).fillna("")
    except Exception as e:
        # CSV 파일에 문제가 있어도 앱 전체가 죽지 않고,
        # 병명 매핑 없이(AI 번역만으로) 계속 동작하도록 함
        st.warning(f"⚠️ 병명 매핑 테이블(diagnosis_mapping.csv)을 읽는 중 오류가 발생하여 매핑 없이 진행합니다: {e}")
        return empty_df


def lookup_english_diagnosis(korean_name: str, mapping_df: pd.DataFrame) -> str:
    """매핑 테이블에서 정확히 일치하는 한글 병명을 찾아 영문명을 반환.
    없으면 빈 문자열을 반환하여, 호출부에서 AI 번역으로 대체하도록 한다."""
    if not korean_name:
        return ""
    match = mapping_df[mapping_df["korean_name"].str.strip() == korean_name.strip()]
    if not match.empty:
        return match.iloc[0]["english_name"]
    # 부분 일치도 시도 (예: "급성 심근경색증(전벽)" 처럼 수식어가 붙은 경우)
    contains = mapping_df[mapping_df["korean_name"].apply(lambda x: x in korean_name or korean_name in x)]
    if not contains.empty:
        return contains.iloc[0]["english_name"]
    return ""


# ---------------------------------------------------------------------------
# 2. Gemini Vision을 통한 진단서 항목 추출
# ---------------------------------------------------------------------------

EXTRACTION_PROMPT = """
당신은 한국 병원에서 발급한 진단서(진단서/소견서) 이미지를 분석하는 전문가입니다.
아래 이미지는 병원마다 서식이 다른 한글 진단서입니다. 서식에 관계없이 다음 항목을
최대한 정확하게 찾아서 순수 JSON 형식으로만 응답하세요. 항목을 찾을 수 없으면
빈 문자열("")로 두세요. 설명, 코드블록 기호(```), 그 외 텍스트는 절대 포함하지 마세요.

{
  "hospital_name": "발급 병원명",
  "patient_name_kor": "환자 성명 (한글)",
  "patient_name_eng": "환자 성명 로마자 표기 시도 (여권 표기가 없다면 추정 표기, 없으면 빈값)",
  "patient_birth_date": "생년월일 (YYYY-MM-DD 형식으로 변환)",
  "patient_gender": "성별 (M/F, 확인 불가시 빈값)",
  "diagnosis_korean": "진단명(병명) 원문 그대로 (한글)",
  "diagnosis_code": "진단서에 명시된 상병코드(KCD/ICD)가 있다면 그대로, 없으면 빈값",
  "onset_date": "발병일 (YYYY-MM-DD)",
  "diagnosis_date": "진단일 (YYYY-MM-DD)",
  "issue_date": "진단서 발급일 (YYYY-MM-DD)",
  "clinical_summary_korean": "치료 경과 및 향후 소견 원문 (한글, 요약하지 말고 원문 최대한 그대로)",
  "doctor_name": "발급 의사 성명",
  "doctor_license_no": "의사 면허번호",
  "hospital_phone": "병원 전화번호 (있는 경우)",
  "hospital_address": "병원 주소 (있는 경우)"
}
"""


def call_gemini_vision_extract(image_bytes: bytes, mime_type: str, api_key: str) -> dict:
    """Gemini API에 이미지를 보내 진단서 항목을 JSON으로 추출한다.
    google-genai SDK 사용 (기존 프로젝트들과 동일한 라이브러리)."""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=api_key)

    response = client.models.generate_content(
        model="gemini-3.6-flash",
        contents=[
            types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
            EXTRACTION_PROMPT,
        ],
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )

    raw_text = response.text.strip()
    # 혹시 모델이 코드블록으로 감싸는 경우를 대비한 방어 코드
    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        raw_text = raw_text.replace("json", "", 1).strip()

    return json.loads(raw_text)


def call_gemini_translate_diagnosis(korean_diagnosis: str, api_key: str) -> str:
    """매핑 테이블에 없는 병명을 AI로 보조 번역. 실제 서비스에서는 이 결과를
    반드시 원무과/의료진이 재확인하도록 UI에서 강조 표시한다."""
    from google import genai

    client = genai.Client(api_key=api_key)
    prompt = (
        "다음은 한국 진단서에 기재된 병명입니다. 국제적으로 통용되는 "
        "영문 의학 진단명(ICD-10 표준 명칭 기준)으로만 답하세요. "
        "다른 설명 없이 영문 병명만 출력하세요.\n\n"
        f"병명: {korean_diagnosis}"
    )
    response = client.models.generate_content(model="gemini-3.6-flash", contents=prompt)
    return response.text.strip()


def call_gemini_translate_free_text(korean_text: str, api_key: str) -> str:
    """치료 경과/소견 등 자유 서술 텍스트를 자연스러운 영문 의무기록 문체로 번역."""
    from google import genai

    client = genai.Client(api_key=api_key)
    prompt = (
        "다음은 한국 진단서에 기재된 치료 경과 및 향후 소견 내용입니다. "
        "해외 제출용 영문 진단서에 들어갈 수 있도록, 의무기록에 쓰이는 "
        "격식있고 간결한 영문으로 번역하세요. 번역문만 출력하고 다른 설명은 "
        "하지 마세요.\n\n"
        f"{korean_text}"
    )
    response = client.models.generate_content(model="gemini-3.6-flash", contents=prompt)
    return response.text.strip()


# ---------------------------------------------------------------------------
# 3. 영문 진단서 Word 문서 생성
# ---------------------------------------------------------------------------

def _add_field_row(table, label: str, value: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = value if value else "-"
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)


def generate_certificate_docx(data: dict, purpose: str) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    title_map = {
        "fit_to_fly": "MEDICAL CERTIFICATE OF FITNESS TO FLY",
        "insurance": "MEDICAL CERTIFICATE FOR INSURANCE CLAIM",
        "general": "MEDICAL CERTIFICATE",
    }

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_map.get(purpose, "MEDICAL CERTIFICATE"))
    run.font.size = Pt(16)
    run.font.bold = True

    hospital_line = doc.add_paragraph()
    hospital_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hospital_line.add_run(data.get("hospital_name", ""))
    hr.font.size = Pt(11)
    hr.font.bold = True

    if data.get("hospital_address") or data.get("hospital_phone"):
        addr_line = doc.add_paragraph()
        addr_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_run = addr_line.add_run(
            f'{data.get("hospital_address", "")}   {data.get("hospital_phone", "")}'.strip()
        )
        addr_run.font.size = Pt(9)

    doc.add_paragraph()

    # 환자 정보 테이블
    patient_table = doc.add_table(rows=0, cols=2)
    patient_table.style = "Table Grid"
    patient_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    patient_table.columns[0].width = Cm(5.0)
    patient_table.columns[1].width = Cm(11.5)

    _add_field_row(patient_table, "Patient Name", data.get("patient_name_eng") or data.get("patient_name_kor", ""))
    _add_field_row(patient_table, "Date of Birth", data.get("patient_birth_date", ""))
    _add_field_row(patient_table, "Gender", data.get("patient_gender", ""))

    doc.add_paragraph()

    # 진단 정보 테이블
    diag_table = doc.add_table(rows=0, cols=2)
    diag_table.style = "Table Grid"
    diag_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    diag_table.columns[0].width = Cm(5.0)
    diag_table.columns[1].width = Cm(11.5)

    diagnosis_line = data.get("diagnosis_english", "")
    if data.get("diagnosis_code"):
        diagnosis_line = f'{diagnosis_line}  (Code: {data.get("diagnosis_code")})'

    _add_field_row(diag_table, "Diagnosis", diagnosis_line)
    _add_field_row(diag_table, "Date of Onset", data.get("onset_date", ""))
    _add_field_row(diag_table, "Date of Diagnosis", data.get("diagnosis_date", ""))
    _add_field_row(diag_table, "Date of Issue", data.get("issue_date", ""))

    doc.add_paragraph()

    # 임상 소견
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("Clinical Summary / Recommendation")
    summary_run.font.bold = True
    summary_run.font.size = Pt(11)

    summary_body = doc.add_paragraph(data.get("clinical_summary_english", ""))
    summary_body.runs[0].font.size = Pt(10.5) if summary_body.runs else None

    # Fit to Fly 전용 문구
    if purpose == "fit_to_fly":
        doc.add_paragraph()
        fit_heading = doc.add_paragraph()
        fit_run = fit_heading.add_run("Fitness to Fly Statement")
        fit_run.font.bold = True
        fit_run.font.size = Pt(11)
        fit_body = doc.add_paragraph(
            "Based on the current clinical condition described above, the patient "
            "is considered ______________ (fit / fit with conditions / not fit) "
            "to travel by air as of the date of this certificate. "
            "[병원에서 해당 사항에 체크 또는 문구 수정 필요]"
        )
        if fit_body.runs:
            fit_body.runs[0].font.size = Pt(10.5)
            fit_body.runs[0].font.italic = True

    doc.add_paragraph()

    # 서명란
    sign_table = doc.add_table(rows=0, cols=2)
    sign_table.columns[0].width = Cm(9.0)
    sign_table.columns[1].width = Cm(7.5)
    row = sign_table.add_row()
    row.cells[0].text = f'Physician: {data.get("doctor_name", "")}'
    row.cells[1].text = "Signature: ______________________"
    row = sign_table.add_row()
    row.cells[0].text = f'License No.: {data.get("doctor_license_no", "")}'
    row.cells[1].text = "Hospital Stamp:"

    doc.add_paragraph()

    # 면책 안내 문구 (모든 출력물에 필수 포함)
    doc.add_paragraph()
    disclaimer_heading = doc.add_paragraph()
    d_run = disclaimer_heading.add_run("Notice")
    d_run.font.bold = True
    d_run.font.size = Pt(9)
    disclaimer_body = doc.add_paragraph(DISCLAIMER_TEXT)
    if disclaimer_body.runs:
        disclaimer_body.runs[0].font.size = Pt(8)
        disclaimer_body.runs[0].font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# 4. Streamlit UI
# ---------------------------------------------------------------------------

def main():
    st.title("🩺 해외 출국용 영문 진단서 변환 도구")
    st.caption(
        "한글 진단서를 업로드하면 AI가 항목을 추출하고, 검토 후 영문 초안(Word)을 생성합니다. "
        "생성된 문서는 반드시 담당 의사의 확인 및 서명·직인 날인 후 제출해야 합니다."
    )

    api_key = st.secrets.get("GEMINI_API_KEY", "") if hasattr(st, "secrets") else ""
    if not api_key:
        api_key = st.text_input("Gemini API Key (테스트용, 배포 시에는 Secrets에 등록하세요)", type="password")

    mapping_df = load_mapping_table()

    # 세션 상태 초기화
    if "extracted" not in st.session_state:
        st.session_state.extracted = None

    st.divider()
    st.subheader("1단계: 용도 선택")
    purpose_label = st.radio("이 영문 진단서는 어떤 용도로 제출하나요?", list(PURPOSE_OPTIONS.values()))
    purpose_key = [k for k, v in PURPOSE_OPTIONS.items() if v == purpose_label][0]

    st.divider()
    st.subheader("2단계: 한글 진단서 업로드")
    uploaded_file = st.file_uploader("진단서 스캔본/사진 업로드 (JPG, PNG, PDF 첫 페이지 권장)", type=["jpg", "jpeg", "png"])

    if uploaded_file is not None:
        st.image(uploaded_file, caption="업로드된 진단서", width=350)

        if st.button("🔍 AI로 항목 추출하기", type="primary", disabled=not api_key):
            with st.spinner("Gemini Vision으로 진단서를 분석하는 중입니다..."):
                try:
                    image_bytes = uploaded_file.getvalue()
                    mime_type = uploaded_file.type or "image/jpeg"
                    extracted = call_gemini_vision_extract(image_bytes, mime_type, api_key)

                    # 병명 영문 매핑 시도
                    korean_diag = extracted.get("diagnosis_korean", "")
                    eng_diag = lookup_english_diagnosis(korean_diag, mapping_df)
                    if not eng_diag and korean_diag:
                        eng_diag = call_gemini_translate_diagnosis(korean_diag, api_key)
                        extracted["diagnosis_mapped_from"] = "AI_translation"
                    else:
                        extracted["diagnosis_mapped_from"] = "mapping_table" if eng_diag else "none"

                    extracted["diagnosis_english"] = eng_diag

                    # 임상 소견 번역
                    if extracted.get("clinical_summary_korean"):
                        extracted["clinical_summary_english"] = call_gemini_translate_free_text(
                            extracted["clinical_summary_korean"], api_key
                        )
                    else:
                        extracted["clinical_summary_english"] = ""

                    st.session_state.extracted = extracted
                    st.success("추출 완료! 아래에서 내용을 확인 및 수정해주세요.")
                except Exception as e:
                    st.error(f"추출 중 오류가 발생했습니다: {e}")

    if st.session_state.extracted:
        st.divider()
        st.subheader("3단계: 추출 결과 검수 (필수)")

        if st.session_state.extracted.get("diagnosis_mapped_from") == "AI_translation":
            st.warning(
                "⚠️ 이 병명은 매핑 테이블에 없어 AI가 직접 번역했습니다. "
                "정확한 진단명인지 반드시 의료진이 재확인해주세요."
            )

        data = st.session_state.extracted
        col1, col2 = st.columns(2)

        with col1:
            data["hospital_name"] = st.text_input("병원명 (Hospital Name)", data.get("hospital_name", ""))
            data["patient_name_eng"] = st.text_input(
                "환자 성명 영문 (여권 표기 기준으로 확인 필요)",
                data.get("patient_name_eng") or data.get("patient_name_kor", ""),
            )
            data["patient_birth_date"] = st.text_input("생년월일 (YYYY-MM-DD)", data.get("patient_birth_date", ""))
            data["patient_gender"] = st.selectbox(
                "성별", ["", "M", "F"],
                index=["", "M", "F"].index(data.get("patient_gender", "")) if data.get("patient_gender") in ["", "M", "F"] else 0,
            )
            data["doctor_name"] = st.text_input("발급 의사 성명", data.get("doctor_name", ""))
            data["doctor_license_no"] = st.text_input("의사 면허번호", data.get("doctor_license_no", ""))

        with col2:
            data["diagnosis_english"] = st.text_input("진단명 영문 (Diagnosis, 확인 필수)", data.get("diagnosis_english", ""))
            data["diagnosis_code"] = st.text_input("진단 코드 (KCD/ICD, 선택)", data.get("diagnosis_code", ""))
            data["onset_date"] = st.text_input("발병일 (YYYY-MM-DD)", data.get("onset_date", ""))
            data["diagnosis_date"] = st.text_input("진단일 (YYYY-MM-DD)", data.get("diagnosis_date", ""))
            data["issue_date"] = st.text_input("발급일 (YYYY-MM-DD)", data.get("issue_date", str(date.today())))
            data["hospital_address"] = st.text_input("병원 주소 (선택)", data.get("hospital_address", ""))
            data["hospital_phone"] = st.text_input("병원 전화번호 (선택)", data.get("hospital_phone", ""))

        data["clinical_summary_english"] = st.text_area(
            "치료 경과 및 향후 소견 (영문, 확인 및 수정 필요)",
            data.get("clinical_summary_english", ""),
            height=150,
        )

        with st.expander("원문(한글) 추출 결과 보기"):
            st.json({
                "diagnosis_korean": data.get("diagnosis_korean", ""),
                "clinical_summary_korean": data.get("clinical_summary_korean", ""),
            })

        st.session_state.extracted = data

        st.divider()
        st.subheader("4단계: 영문 진단서 생성")

        if st.button("📄 영문 진단서 Word 파일 생성", type="primary"):
            docx_bytes = generate_certificate_docx(data, purpose_key)
            filename = f"english_medical_certificate_{data.get('patient_name_eng', 'patient').replace(' ', '_')}.docx"
            st.download_button(
                label="⬇️ 다운로드",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.info(
                "생성된 문서는 초안입니다. 담당 의사의 검토·서명과 병원 직인 날인 후 "
                "제출해주세요."
            )


if __name__ == "__main__":
    main()
